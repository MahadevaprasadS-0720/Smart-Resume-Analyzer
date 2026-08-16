import io
from typing import Tuple, Dict, Any
import PyPDF2
import docx
from app.utils.text_cleaner import clean_text


def parse_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from PDF bytes using PyPDF2.
    Returns extracted text and total page count.
    """
    pdf_file = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(pdf_file)
    extracted_text = []
    page_count = len(reader.pages)

    for page in reader.pages:
        text = page.extract_text()
        if text:
            extracted_text.append(text)

    full_text = "\n".join(extracted_text)
    return clean_text(full_text), page_count


def parse_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from DOCX bytes using python-docx.
    Returns extracted text and estimated page count.
    """
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    full_text = []

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    cleaned = clean_text("\n".join(full_text))
    # Approximate page count based on 400 words per page
    words = len(cleaned.split())
    page_count = max(1, (words + 399) // 400)
    return cleaned, page_count


def parse_text_stream(file_bytes: bytes, filename: str) -> Tuple[str, int]:
    """
    Routes file bytes to the appropriate parser based on file extension.
    """
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
        return parse_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
        cleaned = clean_text(text)
        words = len(cleaned.split())
        page_count = max(1, (words + 399) // 400)
        return cleaned, page_count
    else:
        # Fallback to UTF-8 decoding
        text = file_bytes.decode("utf-8", errors="ignore")
        return clean_text(text), 1


def parse_resume_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Parses resume binary stream and computes text and basic metadata.
    """
    text, page_count = parse_text_stream(file_bytes, filename)
    words = text.split()
    return {
        "text": text,
        "filename": filename,
        "page_count": page_count,
        "word_count": len(words),
        "character_count": len(text),
    }
