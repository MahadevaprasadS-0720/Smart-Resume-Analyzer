import io
import re
from typing import Tuple, Dict, Any
import PyPDF2
import docx


def clean_extracted_text(text: str) -> str:
    """
    Cleans, normalizes, and decodes raw extracted text from resumes.
    - Fixes unicode ligatures (e.g. fi, fl).
    - Removes unreadable non-printable characters.
    - Standardizes bullet points and excessive white spaces.
    """
    if not text:
        return ""

    # Replace common PDF ligatures
    ligatures = {
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
        "\ufb05": "ft",
        "\ufb06": "st",
        "\xa0": " ",
        "\u200b": "",
        "\u200e": "",
        "\u200f": "",
    }
    for orig, rep in ligatures.items():
        text = text.replace(orig, rep)

    # Standardize bullet symbols
    text = re.sub(r"[•●▪■◆★►✓✔*·\t]+", " ", text)

    # Replace carriage returns with standard newlines
    text = re.sub(r"\r\n|\r", "\n", text)

    # Collapse multiple consecutive spaces to a single space
    text = re.sub(r"[ \t]+", " ", text)

    # Limit consecutive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from PDF bytes using PyPDF2 with page tracking and robust fallback.
    Returns: (cleaned_text, page_count)
    """
    pdf_stream = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(pdf_stream)
    page_count = len(reader.pages)
    extracted_pages = []

    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                extracted_pages.append(page_text.strip())
        except Exception:
            continue

    raw_text = "\n\n".join(extracted_pages)
    cleaned = clean_extracted_text(raw_text)
    return cleaned, max(1, page_count)


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts structured text from DOCX bytes using python-docx.
    Processes paragraphs, lists, and table rows.
    Returns: (cleaned_text, estimated_page_count)
    """
    docx_stream = io.BytesIO(file_bytes)
    doc = docx.Document(docx_stream)
    chunks = []

    # 1. Extract paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            chunks.append(txt)

    # 2. Extract tables (resumes often use tables for layouts)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_cells:
                # Remove duplicate adjacent cells caused by merged table cells
                unique_cells = []
                for cell in row_cells:
                    if not unique_cells or cell != unique_cells[-1]:
                        unique_cells.append(cell)
                chunks.append(" | ".join(unique_cells))

    raw_text = "\n".join(chunks)
    cleaned = clean_extracted_text(raw_text)

    # Approximate page count (approx 400 words per page)
    words = len(cleaned.split())
    page_count = max(1, (words + 399) // 400)
    return cleaned, page_count


def extract_text_from_txt(file_bytes: bytes) -> Tuple[str, int]:
    """
    Decodes plain text bytes with fallback encodings.
    """
    encodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
    for enc in encodings:
        try:
            raw_text = file_bytes.decode(enc)
            cleaned = clean_extracted_text(raw_text)
            words = len(cleaned.split())
            page_count = max(1, (words + 399) // 400)
            return cleaned, page_count
        except UnicodeDecodeError:
            continue
    # Force decoding with replacement
    raw_text = file_bytes.decode("utf-8", errors="replace")
    cleaned = clean_extracted_text(raw_text)
    return cleaned, 1


def parse_resume(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Unified parsing router for uploaded resume files.
    Determines parser from file extension and returns text + metadata.
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        text, pages = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx") or lower.endswith(".doc"):
        text, pages = extract_text_from_docx(file_bytes)
    elif lower.endswith(".txt"):
        text, pages = extract_text_from_txt(file_bytes)
    else:
        # Fallback generic text parser
        text, pages = extract_text_from_txt(file_bytes)

    words = text.split()
    return {
        "text": text,
        "filename": filename,
        "page_count": pages,
        "word_count": len(words),
        "character_count": len(text),
    }
